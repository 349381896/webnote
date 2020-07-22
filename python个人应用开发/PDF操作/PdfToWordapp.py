from pdfminer.pdfparser import PDFParser, PDFDocument#文档分析器、文档对象
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter#资源管理器、解释器
from pdfminer.layout import LAParams#参数分析仪
from pdfminer.converter import PDFPageAggregator#聚合器
from pdfminer.pdfinterp import PDFTextExtractionNotAllowed
from docx import Document
 
document = Document()
 
 
def parse(file,outfile):
    # rb以二进制读模式打开本地pdf文件

    fn = open(file.replace('\\', '/'),'rb')
    # 创建一个pdf文档分析器
    parser = PDFParser(fn)
    # 创建一个PDF文档
    doc = PDFDocument()
    # 连接分析器 与文档对象
    parser.set_document(doc)
    doc.set_parser(parser)
 
    
    doc.initialize("")
    # 检测文档是否提供txt转换，不提供就忽略
    if not doc.is_extractable:
        raise PDFTextExtractionNotAllowed
 
    else:
        # 创建PDf资源管理器
        resource = PDFResourceManager()
        # 创建一个PDF参数分析器
        laparams = LAParams()
        # 创建聚合器,用于读取文档的对象
        device = PDFPageAggregator(resource,laparams=laparams)
        # 创建解释器，对文档编码，解释成Python能够识别的格式
        interpreter = PDFPageInterpreter(resource,device)
        # 循环遍历列表，每次处理一页的内容
        # doc.get_pages() 获取page列表
        for page in doc.get_pages():
            # 利用解释器的process_page()方法解析读取单独页数
            interpreter.process_page(page)
            # 使用聚合器get_result()方法获取内容
            layout = device.get_result()
            # 这里layout是一个LTPage对象,里面存放着这个page解析出的各种对象
            for out in layout:
                # 判断是否含有get_text()方法，获取我们想要的文字
                if hasattr(out,"get_text"):
                    # print(out.get_text(), type(out.get_text()))
                    content = out.get_text().replace(u'\xa0', u' ')  # 将'\xa0'替换成u' '空格，这个\xa0就是&nbps空格
                    # with open('test.txt','a') as f:
                    #     f.write(out.get_text().replace(u'\xa0', u' ')+'\n')
                    document.add_paragraph(
                        content, style='ListBullet'    # 添加段落，样式为unordered list类型
                    )
                document.save(outfile.replace('\\', '/'))  # 保存这个文档
        return 1
 
if __name__ == '__main__':
    parse()